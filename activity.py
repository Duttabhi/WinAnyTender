import os
import re
import shutil
import pythoncom
from docx import Document
from docx.shared import Inches
import win32com.client as win32
from docxprocess import grab_image_and_text
from docx.enum.text import WD_ALIGN_PARAGRAPH
from win32com.client import constants as wcnst
# LLMs
from chatgpt import chatGPTApi
from bardapi import Bard
os.environ['_BARD_API_KEY']="iG3ZsIdFlkyyf13Un4cIpPE94TLrbWIXzIDbVYjr6ZR9oxmZNRuwLdAD3rs8xl8Ixi3A."

class Process:
    def __init__(self, request, request_time):
        # Get the text data from the request
        self.business_option = request.form.get('businessOption')
        self.company_name = request.form.get('companyName')
        self.company_abn = request.form.get('abn')
        self.company_acn = request.form.get('acn')
        self.company_estd = request.form.get('estd')
        self.company_addr = request.form.get('addr')
        self.company_license = request.form.get('license')
        self.company_owner = request.form.get('owner')
        self.company_additional_data = request.form.get('additionalData')
        self.awards = request.form.get('awards')
        self.phone = request.form.get('phone')
        self.email = request.form.get('email')
        self.companyService = request.form.get('companyService')
        self.numEmployees = request.form.get('numEmployees')
        self.request_time = request_time
        self.documents_folder = "documents\\" + request_time
        self.templates_folder = "templates\\"
        self.path_prefix = "C:\\Users\\abhis\\Documents\\Flask Projects\\Profiler"

        # Process the work pictures if it's included in the request
        id = 0
        pictures = request.files.getlist('workPic')
        for picture in pictures:
            id = id + 1
            try:
                if not os.path.exists(self.documents_folder):
                    os.makedirs(self.documents_folder)
                output_file_path = os.path.join(self.documents_folder, "workPic_"+str(id)+".jpg")
                picture.save(output_file_path)
                # Process the image here, e.g., save it to a folder or perform some image processing
            except Exception as e:
                print("Error processing image file" + str(e))
        self.workPicCount = id   

        # Process the document prev work if it's included in the request
        id = 0
        prevdocs = request.files.getlist('prevWork')
        for doc in prevdocs:
            id = id + 1
            try:
                if not os.path.exists(self.documents_folder):
                    os.makedirs(self.documents_folder)
                output_file_path = os.path.join(self.documents_folder, "prevWork_"+str(id)+".docx")
                doc.save(output_file_path)
            except Exception as e:
                print("Error processing previous work document file" + e)  
        self.prevWorkDocCount = id

        # Process the document owner docs if it's included in the request
        id = 0
        owndocs = request.files.getlist('ownerDocs')
        for doc in owndocs:
            id = id + 1
            try:
                if not os.path.exists(self.documents_folder):
                    os.makedirs(self.documents_folder)
                output_file_path = os.path.join(self.documents_folder, "ownerDocs_"+str(id)+".docx")
                doc.save(output_file_path)
            except Exception as e:
                print("Error processing owner document file" + str(e))
        self.ownDocCount = id

        # Process the document key emp details if it's included in the request
        self.employees = []
        id = 0
        print("Processing number of employee details count: " + str(self.numEmployees))
        while id < int(self.numEmployees):
            employee_data = {}
            employee_data['name'] = request.form.get(f'emp[{id}][eName]')  
            employee_data['designation'] = request.form.get(f'emp[{id}][eDesignation]')       
            cv_file = request.files.get(f'emp[{id}][eCv]')
            id = id + 1
            
            if cv_file:
                output_file_path = os.path.join(self.documents_folder, "employee_"+str(id)+".docx")
                cv_file.save(output_file_path)
                employee_data['cv'] = "employee_"+str(id)+".docx"
            # Append the employee data
            self.employees.append(employee_data)
        # Print Employee Info 
        print(self.employees)

    # Grab text from the text document
    def fetch_document_data(self, file_path):
        try:
            doc = Document(file_path)
            document_text = ""
            for para in doc.paragraphs:
                document_text += para.text + '\n'
            print("File Path: " + file_path)
            return document_text
        except Exception as e:
            print("Error fetching document data" + str(e))

    # Replace Image
    def find_replace_image(self, doc_path1, find_text, replace_img_path, workpiccnt):
        try:
            print("Image Path: " + replace_img_path)
            print("Number of images to be replaced:" + str(workpiccnt))
            doc1 = Document(doc_path1)
            for paragraph in doc1.paragraphs:
                if find_text in paragraph.text:
                    paragraph.clear()
                    for idx in range(workpiccnt):
                        print("Adding image with file name: " + f"workPic_{idx+1}.jpg")
                        run = paragraph.add_run()
                        run.add_picture(os.path.join(replace_img_path, f"workPic_{idx+1}.jpg"), width=Inches(4))
                    break
            doc1.save(doc_path1)
        except Exception as e:
            print("Error attaching image files" + str(e)) 

    # Replace Image
    def find_replace_text_image(self, doc_path2, find_text, replace_text, replace_img):
        # Initialize the COM library if not already initialized
        pythoncom.CoInitialize()
        # Open an instance of Word
        word_app = win32.Dispatch("Word.Application")
        # word_app.Visible = False  # Set to True if you want to see Word in action
        print("Find and replace image and text started with count of text to be replaced: " + str(len(replace_text)))
        try:
            doc2 = word_app.Documents.Open(doc_path2)
            numberofimgs = len(replace_img)
            numberofpara = len(replace_text)
            for paragraph in doc2.Paragraphs:
                if find_text in paragraph.Range.Text:
                    # Find the keyword and replace it with an empty string to remove it
                    paragraph.Range.Text.replace(find_text, "")
                    for paraid in range(numberofpara):
                        # print(str(replace_text[paraid]))
                        if paraid == 0:
                            paragraph.Range.Text = str(replace_text[paraid])
                        else:
                            paragraph.Range.InsertAfter(str(replace_text[paraid]))

                        # paragraph.Alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.Alignment = wcnst.wdAlignParagraphLeft 
                        paragraph.Range.InsertAfter(f"\n")
                        newPara = paragraph.Range.Paragraphs.Add()
                        newPara.Range.Font.Bold = False
                        newPara.Range.Font.Size = 12
                        newPara.Range.InsertAfter(f"\n")
                        newerPara = newPara.Range.Paragraphs.Add()                                              
                            
                        
                        if paraid < numberofimgs:
                            newimgpath = self.path_prefix + "\\" + self.documents_folder + "\\" + replace_img[paraid]
                            print("Replacing image: " + newimgpath)                       
                            image_range = doc2.Content.Duplicate
                            image_shape = image_range.InlineShapes.AddPicture(newimgpath, LinkToFile=False, SaveWithDocument=True, Range=newerPara.Range)
                            # Set the size of the image
                            image_shape.Width = 200
                            image_shape.Height = 120
                            image_range.Paragraphs.Add()
                            image_range.InsertAfter(f"\n")    
                    break
            # Save and close the document
            doc2.Save()
            doc2.Close()


            print("Find and replace image and text ended")
            return True
        except Exception as e:
            print("Error attaching image files" + str(e))
            return False
        finally:
            # Quit Word application
            word_app.Quit()
            # Uninitialize the COM library
            pythoncom.CoUninitialize() 

    # Replace text function
    def find_replace_text(self, doc_path3, find_text, replace_text, isAlignReq = True):

        if replace_text == "" or replace_text.isspace():
            return

        # Initialize the COM library if not already initialized
        pythoncom.CoInitialize()

        # Open an instance of Word
        word_app = win32.Dispatch("Word.Application")
        word_app.Visible = False  # Set to True if you want to see Word in action
        try:
            # Open the input document
            doc3 = word_app.Documents.Open(doc_path3)

            # Get all paragraphs in the document
            paragraphs = doc3.Paragraphs

            # Flag to check if the target word was replaced
            word_replaced = False

            # Find the target word in the document
            for paragraph in paragraphs:
                if find_text in paragraph.Range.Text:
                    # Replace the target word with the new paragraph text
                    paragraph.Range.Text = paragraph.Range.Text.replace(find_text, replace_text)
                    if isAlignReq == True:
                        # paragraph.Alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.Alignment = wcnst.wdAlignParagraphLeft
                        paragraph.Range.Font.Bold = False
                        paragraph.Range.Font.Size = 12
                    word_replaced = True
                    print(f"Target word: {find_text} is found and replaced")
                    # break
            # If the target word was not found, notify the user
            if not word_replaced:
                print(f"Target word '{find_text}' not found in the document.")

            # Save and close the document
            doc3.Save()
            doc3.Close()

        except Exception as e:
            print(f"Error: {e}")

        finally:
            # Quit Word application
            word_app.Quit()
            # Uninitialize the COM library
            pythoncom.CoUninitialize()            

    # Copy template_{id} from template\ folder to documents_folder with result_{request_time}.docx
    def copy_template(self, template_id):
        source_path = os.path.join(self.templates_folder, "template_"+template_id+".docx")
        destination_path = os.path.join(self.documents_folder, "result_"+self.request_time+".docx")
        try:
            shutil.copy2(source_path, destination_path)
            print(f"File successfully copied from '{source_path}' to '{destination_path}'.")
        except FileNotFoundError:
            print(f"File not found in '{source_path}'.")
        except PermissionError:
            print(f"Permission denied. Unable to copy '{source_path}'.")  

    # Just a test function
    def process_test(self):
        # Implement your logic here to process the text data and generate a response
        # For example, you can use conditionals, apply some algorithms, etc.
        # Replace this with your own custom logic based on your application requirements
        print("Processing info about: " + self.company_name)
        return 'Request successfully processed at backend: ' + self.company_name

    # Extract in=between text if required
    def extract_in_between_paragraphs(self, document):
        paragraphs = document.strip().split('\n')
        # in_between_paragraphs = paragraphs[1:-1]
        # return_text = '\n'.join(in_between_paragraphs)
        return_text = ""
        for para in paragraphs:
            if(len(para) > 5):
                return_text = return_text + "\n" + para
        # print(return_text)
        if len(return_text) < 10:
            return_text = ''.join(document)
        return return_text

    # handle LLM model
    def call_llm_api(self, query):
        try:
            print("Calling LLM Model")
            # response = self.extract_in_between_paragraphs(Bard().get_answer(query)['content'])
            response = self.extract_in_between_paragraphs(chatGPTApi(query))
            print("Success calling LLM MOdel")
            # print(response)
            return response
        except Exception as e:
            print("Exception calling LLM: " + str(e))
            return ""

    # Find red label
    def beutify_doc(self, doc_path4):
        # Initialize the COM library if not already initialized
        pythoncom.CoInitialize()
        # Initialize Word application
        word_app = win32.Dispatch("Word.Application")
        word_app.Visible = False
        try:
            # Load the existing document
            doc4 = word_app.Documents.Open(doc_path4)
            # Define a regular expression pattern to match text between double asterisks
            pattern = r"\*\*(.*?)\*\*"
            # Iterate through all matches in the document
            for match in re.finditer(pattern, doc4.Content.Text):
                text_to_format = match.group(1)  # Get the text between double asterisks
                # Create a Range object for the matched text
                range = doc4.Content.Duplicate
                range.Find.Execute(FindText=match.group(0))
                # Apply formatting (bold and red font color)
                range.Font.Bold = True
                range.Font.Color = 255  # Red color
            # Save the modified document (optional)
            doc4.Save()
            # Close the document and quit Word application
            doc4.Close()
        except Exception as e:
            print("Error while beutifying" + str(e))
        finally:
            word_app.Quit()
            # Uninitialize the COM library
            pythoncom.CoUninitialize()

    def beutify2(self, doc_path):
        # Initialize the COM library if not already initialized
        pythoncom.CoInitialize()

        # Open an instance of Word
        word_app = win32.Dispatch("Word.Application")
        word_app.Visible = False  # Set to True if you want to see Word in action

        try:
            # Open the input document
            doc = word_app.Documents.Open(doc_path)

            # Get all paragraphs in the document
            paragraphs = doc.Paragraphs

            # Find the target word in the document
            for paragraph in paragraphs:
                # paragraph.Alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.Alignment = wcnst.wdAlignParagraphLeft

            # Save and close the document
            doc.Save()
            doc.Close()

        except Exception as e:
            print(f"Error: {e}")

        finally:
            # Quit Word application
            word_app.Quit()
            # Uninitialize the COM library
            pythoncom.CoUninitialize()

    # Process result_{request_time}.docx in documents_folder
    def overwrite_data(self):
        print("overwrite_data() function called...")
        docx_path = os.path.join(self.path_prefix, self.documents_folder, "result_"+self.request_time+".docx")
        try:
            prevWork_path = os.path.join(self.path_prefix, self.documents_folder, "prevWork_1.docx") # default for now only one document
            img_save_loc = os.path.join(self.path_prefix, self.documents_folder) + "\\"     
            # Important to have previou work info and also append the additional data for company work info
            experience_text = ""
            if os.path.exists(prevWork_path):
                experience_text = self.fetch_document_data(prevWork_path)
            experience_text = experience_text + self.company_additional_data

            self.find_replace_text(docx_path, "__businessName__", self.company_name, False)
            self.find_replace_text(docx_path, "__businessStructure__", self.business_option)
            self.find_replace_text(docx_path, "__abn__", self.company_abn)
            self.find_replace_text(docx_path, "__acn__", self.company_acn)
            self.find_replace_text(docx_path, "__businessLocation__", self.company_addr)
            self.find_replace_text(docx_path, "__established__", self.company_estd)
            self.find_replace_text(docx_path, "__owner__", self.company_owner)
            self.find_replace_text(docx_path, "__phone__", self.phone, False)
            self.find_replace_text(docx_path, "__email__", self.email, False)
            self.find_replace_text(docx_path, "__companyService__", self.companyService)
            self.find_replace_text(docx_path, "__awardAndAchivements__", self.awards)   
        
            # TODO: __vision__
            try:
                find_text = "__vision__"
                query = f"Write a brief pragraph about the vision of {self.company_name} company working in {self.business_option} sector and providing {self.companyService} service."
                replace_text = self.call_llm_api(query)
                # print("vision text:" + replace_text)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))
            
            # TODO: __mission__
            try:
                find_text = "__mission__"
                query = f"Write a brief pragraph about the mission of {self.company_name} company working in {self.business_option} sector and providing {self.companyService} service."
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))        

            # TODO: __background__
            try:
                find_text = "__background__"
                query = f"Write a brief paragraph about the background of {self.company_name} company working in {self.business_option} sector and providing {self.companyService} service. The company was established in {self.company_estd} and is located in {self.company_addr}."
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __experience__
            try:
                find_text = "__experience__"
                query = f"Write a brief paragraph about the experience of {self.company_name} company working in {self.business_option} and established in {self.company_estd}. The company has experience as mentioned in following text: {experience_text}."
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __healthAndSafety__
            try:
                find_text = "__healthAndSafety__"
                query = f"Write 2 brief paragraph about the health and safety measure taken by company {self.company_name} working in {self.business_option} and established in {self.company_estd}, using following reference text as follows: {experience_text}."
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __environment_management__
            try:
                find_text = "__environmentManagement__"
                query = f"Write 2 brief paragraphs about how {self.company_name} has made its environmental mangement plan working in {self.business_option} domain."
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __industrialRelations__
            try:
                find_text = "__industrialRelations__"
                query = f"Write 2 brief paragraph about the industrial mangement and safety measures taken by the company {self.company_name} working in {self.business_option} and providing {self.companyService} service."
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __communityEngagementAndConsultations__
            try:
                find_text = "__communityEngagementAndConsultations__"
                query = f"Write 2 brief paragraph about the community engagement and cosultations by the company {self.company_name} working in {self.business_option} and established in {self.company_estd}, using following reference text: {experience_text}"
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __managementSystems__
            try:
                find_text = "__managementSystems__"
                query = f"Write 2 brief paragraph about the management systems of the company {self.company_name} working in {self.business_option} and established in {self.company_estd}, using following reference text: {experience_text}"
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __recentProjects__
            try:
                find_text = "__recentProjects__"
                # Load the source document contents (img and text) alternatively
                gtext, gimages = grab_image_and_text(prevWork_path, img_save_loc, 'RecentProjects')
                self.find_replace_text_image(docx_path, find_text, gtext, gimages)
            except Exception as e:
                print("Exception:" + str(e))
            

            # TODO: __projectApproachStrategy__
            try:
                find_text = "__projectApproachStrategy__"
                query = f"Write a brief paragraph about the project approach strategy of the company {self.company_name} working in {self.business_option} and established in {self.company_estd}, using following reference text: {experience_text}"
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __ownerInfo__
            find_text = "__ownerInfo__"
            try:
                owner_file_path = os.path.join(self.documents_folder, "ownerDocs_1.docx") # default for now only one document
                gtext2, gimages2 = grab_image_and_text(owner_file_path, img_save_loc, 'OwnerInfo')
                replace_text_list = []
                info_text = ' '.join(gtext2)
                query = f"Write a short resume of 100 words with About, Skills and Experience sections by referring following text: {info_text}"
                replace_text = self.call_llm_api(query)
                replace_text_list.append(replace_text)
                # It will capture only the first image, please make sure only one image
                self.find_replace_text_image(docx_path, find_text, replace_text_list, gimages2)
            except Exception as e:
                print("Exception:" + str(e))
        

            # TODO: __keyPeople__
            try:
                find_text = "__keyPeople__"
                num_emp = int(self.numEmployees)
                print("Number of employees in current context: " + str(num_emp))
                for idx in range(num_emp):
                    print("Preparing key people stats...")
                    key_people_path = os.path.join(self.documents_folder, f"employee_{idx+1}.docx")
                    # Get the contents of the source document
                    print(key_people_path)
                    # source_content = "" + self.fetch_document_data(key_people_path)
                    gtext3, gimages3 = grab_image_and_text(key_people_path, img_save_loc, 'KeyPeople')
                    replace_text_list = []
                    info_text = ' '.join(gtext3)
                    query = f"Write a short resume of 100 words with About, Skills and Experience sections by referring following text: {info_text}"
                    replace_text = self.call_llm_api(query)                
                    if idx+1 < int(self.numEmployees):
                        replace_text = replace_text + "\n__keyPeople__"
                    replace_text_list.append(replace_text)
                    # It will capture only the first image, please make sure only one image
                    self.find_replace_text_image(docx_path, find_text, replace_text_list, gimages3)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: __businessCommitment__
            try:
                find_text = "__businessCommitment__"
                query = f"Write a brief paragraph about the business commitment of the company {self.company_name} working in {self.business_option} and established in {self.company_estd}."
                replace_text = self.call_llm_api(query)
                self.find_replace_text(docx_path, find_text, replace_text)
            except Exception as e:
                print("Exception:" + str(e))

            # TODO: Replace images __imagesAttached__
            try:
                find_text = "__imagesAttached__"
                workpiccnt = int(self.workPicCount)
                workPicPath = os.path.join(self.path_prefix, self.documents_folder)
                self.find_replace_image(docx_path, find_text, workPicPath, workpiccnt)
            except Exception as e:
                print("Exception:" + str(e))        

            # self.beutify_doc(docx_path)
            self.beutify2(docx_path)

            print("Returning updated document...")
        except Exception as err:
            print("Some Execption occured:" + str(err))
        finally:
            return docx_path
